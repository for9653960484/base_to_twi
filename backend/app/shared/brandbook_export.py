"""Формирование DOCX технологической карты по шаблону брендбука."""

from __future__ import annotations

import io
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.core.config import settings
from app.shared.schemas import MaintenanceType

MAINTENANCE_LABELS_RU: dict[str, str] = {
    MaintenanceType.ANNUAL.value: "Годовое ТО",
    MaintenanceType.SEMI_ANNUAL.value: "Полугодовое ТО",
    MaintenanceType.QUARTERLY.value: "Квартальное ТО",
    MaintenanceType.MONTHLY.value: "Месячное ТО",
    MaintenanceType.WEEKLY.value: "Недельное ТО",
    MaintenanceType.DAILY.value: "Ежедневное ТО",
}

BRAND_PRIMARY = RGBColor(0x1A, 0x56, 0x8E)
REL_TEMPLATE_PATH = "templates/brandbook/tech_card_ru.docx"


def default_template_path() -> Path:
    return Path(settings.storage_local_path_resolved) / REL_TEMPLATE_PATH


def ensure_default_tech_card_template() -> Path:
    path = default_template_path()
    if path.is_file():
        return path
    path.parent.mkdir(parents=True, exist_ok=True)
    _write_template_file(path)
    return path


def _write_template_file(path: Path) -> None:
    """Корпоративный шаблон брендбука (базовая вёрстка Base To)."""
    doc = Document()
    title = doc.add_heading("Технологическая карта", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = BRAND_PRIMARY

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta.add_run("Оборудование: ").bold = True
    meta.add_run("{{EQUIPMENT_NAME}}")
    meta.add_run("\nВид ТО: ").bold = True
    meta.add_run("{{MAINTENANCE_TYPE}}")
    meta.add_run("\nНаименование: ").bold = True
    meta.add_run("{{TITLE}}")
    meta.add_run("\nДата: ").bold = True
    meta.add_run("{{GENERATED_AT}}")

    doc.add_heading("Перечень работ", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["№", "Описание работ", "Инструменты", "Безопасность", "Контроль"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    doc.add_paragraph("{{WORK_ITEMS_TABLE}}")
    doc.add_paragraph("")
    footer = doc.add_paragraph("Документ сформирован системой Base To")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.save(str(path))


def _render_from_template(
    template_path: Path,
    *,
    equipment_name: str,
    title: str,
    maintenance_type: str,
    work_items: list[dict[str, Any]],
) -> bytes:
    from docx import Document as DocxReader

    doc = DocxReader(str(template_path))
    replacements = {
        "{{EQUIPMENT_NAME}}": equipment_name,
        "{{MAINTENANCE_TYPE}}": MAINTENANCE_LABELS_RU.get(maintenance_type, maintenance_type),
        "{{TITLE}}": title,
        "{{GENERATED_AT}}": datetime.now(timezone.utc).strftime("%d.%m.%Y"),
        "{{WORK_ITEMS_TABLE}}": "",
    }

    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    # Добавляем строки работ во вторую таблицу или в первую таблицу после заголовка
    if doc.tables:
        table = doc.tables[0]
        for item in work_items:
            tools = ", ".join(item.get("tools") or []) or "—"
            safety = "; ".join(item.get("safety") or []) or "—"
            control = item.get("control_params") or {}
            control_text = "; ".join(f"{k}: {v}" for k, v in control.items()) if control else "—"
            row = table.add_row()
            row.cells[0].text = str(item.get("order", ""))
            row.cells[1].text = str(item.get("description", ""))
            row.cells[2].text = tools
            row.cells[3].text = safety
            row.cells[4].text = control_text

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_tech_card_docx(
    template_path: Path,
    *,
    equipment_name: str,
    title: str,
    maintenance_type: str,
    work_items: list[dict[str, Any]],
) -> bytes:
    path = template_path if template_path.is_file() else ensure_default_tech_card_template()
    return _render_from_template(
        path,
        equipment_name=equipment_name,
        title=title,
        maintenance_type=maintenance_type,
        work_items=work_items,
    )
